#!/usr/bin/env python3
"""Steuerleitfaden Network Marketing (MONAT) — aus Call-Transkript erstellt."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "/Users/patrycjakaczorowska/Downloads/Vorlagen/claude-workspace-vorlage/outputs/steuerleitfaden-network-marketing.docx"

# Farben
GOLD    = RGBColor(0xC9, 0xA0, 0x5A)   # Akzentgold — für Designer anpassbar
DUNKEL  = RGBColor(0x1A, 0x1A, 0x2E)   # Fast-Schwarz
MITTEL  = RGBColor(0x3A, 0x3A, 0x5C)   # Mittleres Dunkelblau
HELL    = RGBColor(0xF7, 0xF5, 0xF0)   # Cremeton (für Hintergründe)
GRAU    = RGBColor(0x6B, 0x6B, 0x6B)   # Grau für Untertitel
WEISS   = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Seitenränder (A4) ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Hilfsfunktionen ────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Setzt Hintergrundfarbe einer Tabellenzelle."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="E8E0D0"):
    """Setzt dezente Rahmen für eine Tabellenzelle."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def remove_table_borders(table):
    """Entfernt alle Tabellenrahmen."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def add_paragraph_with_style(doc, text, font_name='Calibri', font_size=11,
                               bold=False, italic=False, color=None,
                               space_before=0, space_after=6,
                               alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # Goldene Linie oben
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '8')
    top.set(qn('w:space'), '4');   top.set(qn('w:color'), 'C9A05A')
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(15)
    run.font.bold  = True
    run.font.color.rgb = DUNKEL
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = MITTEL
    return p

def add_body(doc, text, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = DUNKEL
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = DUNKEL
    return p

def add_hint_box(doc, title, text, bg='F5F0E8', border_color='C9A05A'):
    """Erstellt eine farbige Info-/Hinweisbox."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    # Linke farbige Linie via Border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '16')
    left.set(qn('w:space'), '0');   left.set(qn('w:color'), border_color)
    tcBorders.append(left)
    tcPr.append(tcBorders)
    cell.width = Cm(14.5)
    # Innenabstand
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top','80'),('bottom','80'),('left','160'),('right','160')]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), val); m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)
    # Titel
    if title:
        tp = cell.paragraphs[0]
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after  = Pt(2)
        tr = tp.add_run(title)
        tr.font.name = 'Calibri'; tr.font.size = Pt(11)
        tr.font.bold = True; tr.font.color.rgb = MITTEL
        # Body
        bp = cell.add_paragraph()
    else:
        bp = cell.paragraphs[0]
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after  = Pt(0)
    br = bp.add_run(text)
    br.font.name = 'Calibri'; br.font.size = Pt(10.5)
    br.font.color.rgb = DUNKEL
    # Abstand nach der Box
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_two_col_table(doc, rows_data, header=None, col_widths=(5.5, 9.0)):
    """Zweispaltige Tabelle: links Bezeichnung, rechts Inhalt."""
    n = len(rows_data) + (1 if header else 0)
    table = doc.add_table(rows=n, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(table)
    row_idx = 0
    if header:
        c0, c1 = table.rows[0].cells
        c0.width = Cm(col_widths[0]); c1.width = Cm(col_widths[1])
        set_cell_bg(c0, '1A1A2E'); set_cell_bg(c1, '1A1A2E')
        for c, t in [(c0, header[0]), (c1, header[1])]:
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(t)
            run.font.name = 'Calibri'; run.font.size = Pt(10)
            run.bold = True; run.font.color.rgb = WEISS
            # Padding
            tc = c._tc; tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side, val in [('top','60'),('bottom','60'),('left','100'),('right','100')]:
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), val); m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
        row_idx = 1
    for i, (left, right) in enumerate(rows_data):
        row = table.rows[row_idx + i]
        c0, c1 = row.cells
        c0.width = Cm(col_widths[0]); c1.width = Cm(col_widths[1])
        bg = 'FAFAF7' if i % 2 == 0 else 'F2EDE5'
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        set_cell_borders(c0, 'E0D8CC'); set_cell_borders(c1, 'E0D8CC')
        for c, t, is_bold in [(c0, left, True), (c1, right, False)]:
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(t)
            run.font.name = 'Calibri'; run.font.size = Pt(10.5)
            run.bold = is_bold; run.font.color.rgb = DUNKEL
            tc = c._tc; tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side, val in [('top','60'),('bottom','60'),('left','100'),('right','100')]:
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), val); m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_checklist_item(doc, text, done=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.3)
    box = '☐' if not done else '☑'
    run = p.add_run(f'{box}  {text}')
    run.font.name = 'Calibri'; run.font.size = Pt(11)
    run.font.color.rgb = DUNKEL
    return p

def spacer(doc, pt=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITELBEREICH
# ══════════════════════════════════════════════════════════════════════════════

# Goldene Kopflinie
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '24')
top.set(qn('w:space'), '1');   top.set(qn('w:color'), 'C9A05A')
pBdr.append(top)
pPr.append(pBdr)
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(16)

add_paragraph_with_style(doc, 'STEUERLEITFADEN', font_name='Calibri', font_size=26,
    bold=True, color=DUNKEL, space_before=0, space_after=4,
    alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph_with_style(doc, 'Network Marketing & MONAT', font_name='Calibri', font_size=16,
    bold=False, color=GOLD, space_before=0, space_after=4,
    alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph_with_style(doc, 'Steuerliche Grundlagen für Partnerinnen in Deutschland',
    font_name='Calibri', font_size=12, color=GRAU,
    space_before=0, space_after=20, alignment=WD_ALIGN_PARAGRAPH.LEFT)

# Expertinnen-Info-Box
add_hint_box(doc,
    'Expertin: Nancy Wackerhagen',
    'Steuerfachwirtin (seit 2003) · Steuerberaterin (seit 2010) · Geschäftsführerin einer Kanzlei mit 6 Niederlassungen und 60 Mitarbeitern.',
    bg='1A1A2E', border_color='C9A05A')

# Goldene Trennlinie
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(4)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
bot.set(qn('w:space'), '1');   bot.set(qn('w:color'), 'C9A05A')
pBdr.append(bot); pPr.append(pBdr)

add_paragraph_with_style(doc,
    'Dieser Leitfaden fasst die wichtigsten Inhalte des Steuer-Calls zusammen. '
    'Er ersetzt keine individuelle steuerrechtliche Beratung.',
    font_name='Calibri', font_size=10, italic=True, color=GRAU,
    space_before=4, space_after=16)

# ══════════════════════════════════════════════════════════════════════════════
# 1. DEIN STATUS ALS UNTERNEHMERIN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '01  Dein Status als Unternehmerin')
add_body(doc,
    'Als MONAT-Partnerin bist du rechtlich gesehen Unternehmerin — unabhängig davon, '
    'ob du ein Gewerbe im Haupt- oder Nebenerwerb angemeldet hast oder ob du als '
    'Kleinunternehmerin oder mit Umsatzsteuerpflicht tätig bist.')

add_two_col_table(doc, [
    ('Nebenerwerb',  'Du hast zusätzlich einen Angestelltenjob. Du bleibst über deinen Arbeitgeber krankenversichert, solange dein Angestelltengehalt höher ist als dein Gewinn aus dem Business.'),
    ('Haupterwerb',  'MONAT ist deine einzige oder primäre Einkommensquelle. Du bist für deine Krankenversicherung selbst zuständig und fällst aus der gesetzlichen KV raus.'),
], header=('Beschäftigungsform', 'Was das bedeutet'))

add_hint_box(doc,
    'Wichtig: Krankenversicherung',
    'Die Krankenkasse prüft immer den Gewinn (nicht den Umsatz!). Übersteigt dein Gewinn aus MONAT dein Angestelltengehalt, wird der Krankenkassenbeitrag auf Basis des Gewinns neu berechnet — auch wenn du formal im Nebenerwerb bist.',
    bg='FFF8EC', border_color='C9A05A')

# ══════════════════════════════════════════════════════════════════════════════
# 2. B2B-UMSATZ UND DIE BESONDERHEIT MIT MONAT
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '02  B2B-Umsatz & MONAT: Die Besonderheit')
add_body(doc,
    'MONAT hat seinen Sitz in Irland. Das ist steuerrechtlich entscheidend, '
    'denn der "Ort der sonstigen Leistung" bei zwei Unternehmern (B2B) liegt immer '
    'dort, wo der Leistungsempfänger seinen Sitz hat.')

add_two_col_table(doc, [
    ('Leistende',         'Du — die MONAT-Partnerin in Deutschland'),
    ('Leistungsempfänger','MONAT Global — Sitz in Irland (EU)'),
    ('Art der Leistung',  'Vermittlung (sonstige Leistung im B2B-Bereich)'),
    ('Ort der Leistung',  'Irland (dort, wo MONAT sitzt)'),
    ('Folge',             'Der Umsatz ist in Deutschland NICHT steuerbar. Die Umsatzsteuerpflicht verlagert sich nach Irland.'),
], header=('Begriff', 'In deinem Fall'))

add_hint_box(doc,
    'Das bedeutet für dich',
    'Deine Rechnung (Provisionsabrechnung) ist immer NETTO — du weist keine deutsche '
    'Umsatzsteuer aus. MONAT führt die Umsatzsteuer in Irland ab. '
    'Das Verfahren dafür heißt: Reverse-Charge-Verfahren (§13b UStG).',
    bg='EEF4FB', border_color='3A3A5C')

# ══════════════════════════════════════════════════════════════════════════════
# 3. UMSATZSTEUER-IDENTIFIKATIONSNUMMER
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '03  Die Umsatzsteuer-Identifikationsnummer (USt-IdNr.)')
add_body(doc,
    'Für EU-weite B2B-Umsätze braucht JEDE Partnerin eine Umsatzsteuer-Identifikationsnummer '
    '(USt-IdNr.) — auch Kleinunternehmerinnen. Ohne diese Nummer ist das Reverse-Charge-Verfahren '
    'nicht anwendbar und die Rechnung gilt als unvollständig.')

add_h2(doc, 'So beantragst du die USt-IdNr.')
add_bullet(doc, 'Neu-Anmeldung: Im Eröffnungsfragebogen des Finanzamts das entsprechende Kreuz anklicken.')
add_bullet(doc, 'Nachträgliche Beantragung: Direkt beim Bundeszentralamt für Steuern in Saarlouis — auch elektronisch möglich (ca. 1 Woche Bearbeitungszeit).')
add_bullet(doc, 'Über den Steuerberater beantragen lassen.')
spacer(doc, 6)

add_hint_box(doc, None,
    '🔑  Deine USt-IdNr. beginnt mit „DE" gefolgt von 8 Stellen (z.B. DE123456789). '
    'Die USt-IdNr. von MONAT (Irland) beginnt mit „IE". Du kannst jede EU-USt-IdNr. '
    'kostenlos online verifizieren: vies.ec.europa.eu',
    bg='F5F0E8', border_color='C9A05A')

add_hint_box(doc,
    'Gilt nur in der EU!',
    'Die Umsatzsteuer-Identifikationsnummer gilt ausschließlich innerhalb der Europäischen Union. '
    'Für Networks mit Sitz in der Schweiz, den USA oder anderen Drittländern gelten andere Regeln — '
    'Einzelfallprüfung mit dem Steuerberater notwendig.',
    bg='FFF0F0', border_color='CC4444')

# ══════════════════════════════════════════════════════════════════════════════
# 4. DAS REVERSE-CHARGE-VERFAHREN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '04  Das Reverse-Charge-Verfahren (§13b UStG)')
add_body(doc,
    'Bei grenzüberschreitenden B2B-Leistungen innerhalb der EU geht die '
    'Steuerschuldnerschaft auf den Leistungsempfänger über — in diesem Fall auf MONAT. '
    'Das nennt man Reverse Charge.')

add_h2(doc, 'Was das praktisch bedeutet')
add_bullet(doc, 'Du stellst deine Rechnung netto aus (0% Umsatzsteuer).')
add_bullet(doc, 'MONAT meldet und zahlt die Umsatzsteuer in Irland.')
add_bullet(doc, 'MONAT kann die Umsatzsteuer als Vorsteuer zurückholen — deshalb ist die Abrechnung immer netto.')
add_bullet(doc, 'Erhältst du 1.000 € Provision, sind das 1.000 € — nicht 1.190 €.')
spacer(doc, 6)

add_h2(doc, 'Pflichtangabe auf jeder Rechnung/Abrechnung')
add_hint_box(doc, None,
    '"Steuerschuldnerschaft geht auf den Leistungsempfänger über" (Hinweis nach §13b UStG)',
    bg='F5F0E8', border_color='C9A05A')

# ══════════════════════════════════════════════════════════════════════════════
# 5. DIE PROVISIONSABRECHNUNG
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '05  Deine Provisionsabrechnung')
add_body(doc,
    'MONAT erstellt für jede Auszahlung automatisch eine Gutschrift — diese gilt als deine '
    'Rechnung. Das MONAT-Template enthält bereits alle Pflichtangaben und ist steuerrechtlich '
    'geprüft. Zu finden: im MONAT-Archiv unter Business-Infos / Onboarding.')

add_h2(doc, 'Pflichtangaben auf der Abrechnung (Checkliste)')
add_two_col_table(doc, [
    ('Leistende & Empfänger',  'Dein Name / Firma + MONAT Ireland'),
    ('Rechnungsnummer',        'Die Nummer der Gutschrift = deine Rechnungsnummer'),
    ('Datum',                  'Ausstellungsdatum der Abrechnung'),
    ('Leistungsbeschreibung',  'Vermittlung von Kundinnen und Partnerinnen'),
    ('Leistungszeitraum',      'z.B. 01.01.2026 – 31.01.2026'),
    ('Nettobetrag',            'Der Provisionsbetrag ohne Umsatzsteuer'),
    ('Umsatzsteuer',           '0 € / 0 % — mit Hinweis auf §13b UStG'),
    ('USt-IdNr. beider Parteien', 'Deine DE... und MONATs IE...-Nummer'),
], header=('Pflichtangabe', 'In deiner Abrechnung'))

add_h2(doc, 'Empfohlene Vorgehensweise')
add_bullet(doc, 'Erstelle eine monatliche Zusammenfassung aller Provisionen (MONAT zahlt wöchentlich + monatlich zum 15.).')
add_bullet(doc, 'Trage den Gesamtbetrag pro Monat in die Excel-Vorlage von MONAT ein.')
add_bullet(doc, 'Speichere alles als PDF ab.')
add_bullet(doc, 'Hefte alle Original-Provisionsabrechnungen dahinter.')
add_bullet(doc, 'Schicke die komplette Zusammenfassung monatlich an deinen Steuerberater.')
spacer(doc, 6)

add_hint_box(doc,
    'Warum monatlich?',
    'Monatliche Abrechnungen geben dir einen klaren Überblick über dein tatsächliches Einkommen — '
    'und helfen dir frühzeitig zu erkennen, ob du Rücklagen für Steuerzahlungen bilden musst.',
    bg='F5F0E8', border_color='C9A05A')

# ══════════════════════════════════════════════════════════════════════════════
# 6. DIE ZUSAMMENFASSENDE MELDUNG (ZM)
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '06  Die Zusammenfassende Meldung (ZM)')
add_body(doc,
    'Zusätzlich zur Umsatzsteuervoranmeldung muss eine sogenannte Zusammenfassende Meldung (ZM) '
    'beim Finanzamt eingereicht werden. Sie dokumentiert deine EU-grenzüberschreitenden '
    'Umsätze gegenüber dem deutschen Fiskus.')

add_two_col_table(doc, [
    ('Was wird gemeldet?',  'Deine Umsatzsteuer-IdNr., MONATs USt-IdNr. und die Gesamtsumme der Umsätze im Meldezeitraum'),
    ('Turnus',              'Meistens quartalsweise — gleicher Turnus wie die Umsatzsteuervoranmeldung'),
    ('Betrag',              'Nur Meldepflicht — KEINE Zahlung. Die Summe in der ZM muss mit der Umsatzsteuervoranmeldung übereinstimmen.'),
    ('Einreichung',         'Zeitgleich mit der Umsatzsteuervoranmeldung'),
    ('Aufwand',             'Wird beim Steuerberater automatisch erstellt und übermittelt. Auch selbst erstellbar.'),
], header=('Frage', 'Antwort'))

# ══════════════════════════════════════════════════════════════════════════════
# 7. KLEINUNTERNEHMERIN ODER UMSATZSTEUERPFLICHTIG?
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '07  Kleinunternehmerin oder Umsatzsteuerpflichtig?')
add_body(doc,
    'Diese Entscheidung triffst du bei der Gewerbeanmeldung oder mit einem Zweizeiler ans Finanzamt. '
    'Für MONAT-Partnerinnen gibt es eine wichtige Besonderheit bei der 25.000-€-Grenze.')

add_h2(doc, 'Kleinunternehmerregelung (§19 UStG) — Die Fakten')
add_two_col_table(doc, [
    ('Grenze laufendes Jahr',  '25.000 € Umsatz im INLAND'),
    ('Grenze Folgejahr',       '100.000 € Gesamtumsatz im laufenden Jahr'),
    ('Maßstab',                'Immer UMSATZ — nicht Gewinn!'),
    ('MONAT-Besonderheit',     'Dein MONAT-Umsatz wird in IRLAND erbracht → zählt NICHT zum deutschen Inlandsumsatz!'),
    ('Folge',                  'Auch mit 50.000 € MONAT-Provision bist du umsatzsteuerrechtlich noch Kleinunternehmerin (wenn kein anderes Inlandsgewerbe über 25.000 €)'),
], header=('Regelung', 'Details'))

spacer(doc, 8)
add_h2(doc, 'Die klare Empfehlung von Nancy:')
add_hint_box(doc, None,
    '→  Direkt umsatzsteuerpflichtig wählen!\n\n'
    'Als Umsatzsteuerpflichtige zahlst du in Deutschland keine Umsatzsteuer auf deine '
    'MONAT-Provisionen (weil der Umsatz in Irland liegt). Gleichzeitig bekommst du aus '
    'ALLEN deinen Betriebsausgaben die Vorsteuer zurück. Als Kleinunternehmerin verlierst '
    'du diesen Vorteil komplett.',
    bg='EEF4FB', border_color='3A3A5C')

add_h2(doc, 'Wechsel ist möglich')
add_bullet(doc, 'Einfacher Zweizeiler ans Finanzamt: "Ich möchte ab sofort der Umsatzsteuerpflicht unterliegen."')
add_bullet(doc, 'Das Finanzamt bestätigt und setzt den Turnus für die Umsatzsteuervoranmeldung (meistens quartalsweise).')
add_bullet(doc, 'Als Umsatzsteuerpflichtige bekommst du pro Quartal eine Erstattung (weil der Umsatz steuerfrei in Irland ist, aber Kosten mit Vorsteuer anfallen).')
spacer(doc, 6)

# ══════════════════════════════════════════════════════════════════════════════
# 8. BETRIEBSAUSGABEN & VORSTEUER
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '08  Was kannst du absetzen? Betriebsausgaben & Vorsteuer')
add_body(doc,
    'Als Unternehmerin kannst du alle Kosten absetzen, die der "Sicherung und Erhaltung '
    'der Einnahmen" dienen. Es gibt keine vorgeschriebene Höhe. Wichtig: '
    'Alle Belege aufheben — digital oder in Papierform!')

add_h2(doc, 'Abzugsfähige Betriebsausgaben (Auswahl)')
add_two_col_table(doc, [
    ('Weiterbildung & Seminare',   'Vollständig absetzbar (Seminare, Bücher, Kurse, Coachings)'),
    ('Telefon & Internet',         'Anteilig oder vollständig absetzbar'),
    ('Hardware',                   'Laptop, Handy, Headset etc. — vollständig absetzbar'),
    ('Reisekosten',                '30 Cent/km (Auto) oder tatsächliche Kosten (Zug, Flug, Taxi) — auch Anfahrt zum Flughafen'),
    ('Events & Business-Trips',    'Hotel, Unterkunft, Verpflegung (als Tagegeld) — vollständig absetzbar'),
    ('Geschenke ans Team',         'z.B. bei Rang-Erreichen — vollständig absetzbar (mit Beleg)'),
    ('Bewirtungskosten',           'Restaurantbesuche mit Teampartnerinnen: 70% absetzbar. Pflicht: Anlass + Namen der bewirteten Personen auf dem Beleg vermerken.'),
    ('Vorführprodukte (MONAT)',    'Demo-Produkte für Veranstaltungen — vollständig absetzbar. Von deutschen Rechnungen 19% Vorsteuer zurückfordern!'),
    ('Home-Party / Hausvorstellung','Einkäufe, Getränke, Dekoration — vollständig absetzbar (mit Dokumentation: Wann? Wer eingeladen?)'),
    ('Steuerberater-Kosten',       'Vollständig absetzbar'),
    ('Branding & Design',          'Logo, Grafiken, Website-Design — vollständig absetzbar'),
    ('Assistenz & Aushilfen',      'Lohnkosten für Mitarbeiterinnen — vollständig absetzbar'),
    ('Büromaterial',               'Druckerpatronen, Papier, Schreibwaren etc.'),
], header=('Kostenart', 'Regelung'))

spacer(doc, 6)
add_hint_box(doc,
    'MONAT-Produkte aus Deutschland: Vorsteuer zurückfordern!',
    'Da MONAT Produkte von einer deutschen Adresse liefert, enthalten die Rechnungen 19% '
    'Umsatzsteuer. Als umsatzsteuerpflichtige Unternehmerin kannst du diese Vorsteuer '
    'vollständig zurückfordern — vorausgesetzt, die Produkte dienen dem Business '
    '(z.B. als Vorführware, nicht für den Eigenverbrauch).',
    bg='EEF4FB', border_color='3A3A5C')

add_hint_box(doc, None,
    '📎  Wer schreibt, der bleibt. Alle Kosten dokumentieren:\n'
    '• Datum des Kaufs\n'
    '• Was wurde gekauft?\n'
    '• Für welchen Business-Zweck?\n'
    '• Wer war dabei (bei Bewirtung)?\n\n'
    'Apps für Belegscans direkt nutzen — am besten monatlich, nicht alles auf Jahresende schieben.',
    bg='F5F0E8', border_color='C9A05A')

# ══════════════════════════════════════════════════════════════════════════════
# 9. STEUERERKLÄRUNGEN & FRISTEN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '09  Steuererklärungen & Fristen im Überblick')

add_two_col_table(doc, [
    ('Einkommensteuererklärung',         'Einmal jährlich — erfasst ALLE Einkünfte (Angestelltenlohn + MONAT-Provisionen + sonstige). Abgabefrist: 31.07. des Folgejahres (ohne steuerliche Vertretung).'),
    ('Umsatzsteuer-Voranmeldung',        'Quartalsweise (bei Gewerbestart, danach nach Umsatzhöhe). Enthält eine Spalte für nicht-deutsche Umsätze (Irland). Ergebnis: Erstattung, da Umsatz steuerfrei, aber Vorsteuer aus Kosten anfällt.'),
    ('Umsatzsteuerjahreserklärung',      'Einmal jährlich — muss jeder Unternehmer abgeben.'),
    ('Zusammenfassende Meldung (ZM)',    'Quartalsweise, zeitgleich mit der Umsatzsteuer-Voranmeldung. Nur Meldung, keine Zahlung.'),
    ('Gewerbesteuererklärung',           'Fällig wenn Gewinn > 24.500 € (Freibetrag Einzelunternehmen). ACHTUNG: Keine Ratenzahlung möglich — volle Summe fällig.'),
], header=('Erklärungsart', 'Details & Besonderheiten'), col_widths=(5.0, 9.5))

spacer(doc, 8)
add_hint_box(doc,
    'Gewerbesteuer: Gewinn, nicht Umsatz!',
    'Die Gewerbesteuer fällt erst an, wenn dein GEWINN (also was nach Abzug aller '
    'Betriebsausgaben übrig bleibt) den Freibetrag von 24.500 € übersteigt. '
    'Umsatz und Gewinn können stark voneinander abweichen.',
    bg='FFF0F0', border_color='CC4444')

# ══════════════════════════════════════════════════════════════════════════════
# 10. RÜCKLAGEN & FINANZPLANUNG
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '10  Rücklagen bilden & finanziell planen')
add_body(doc,
    'Als Unternehmerin gibt es keine automatischen Steuerabzüge. Du trägst die volle '
    'Verantwortung — das ist dein größter Vorteil und deine wichtigste Pflicht.')

add_h2(doc, 'Die Daumenregel')
add_hint_box(doc, None,
    '💡  Lege ca. 1/3 deines Gewinns als Steuerrücklage zurück. '
    'Das schützt vor unangenehmen Überraschungen bei der Jahressteuererklärung.',
    bg='F5F0E8', border_color='C9A05A')

add_h2(doc, 'Praktische Maßnahmen')
add_bullet(doc, 'Monatlicher Kassensturz: Einnahmen vs. Ausgaben — wie viel bleibt wirklich übrig?')
add_bullet(doc, 'Vorauszahlungen anpassen: Steuervorauszahlungen können jederzeit nach oben oder unten angepasst werden — beim Finanzamt beantragen.')
add_bullet(doc, 'Verlustjahre nutzen: Verluste können ins Folgejahr übertragen oder ins Vorjahr zurückgetragen werden.')
add_bullet(doc, 'Krankenversicherung anpassen: Die KK passt Beiträge nach letztem Steuerbescheid an — proaktiv mit der KK sprechen.')
add_bullet(doc, 'Buchhaltungs-App nutzen: Belege direkt scannen und digital archivieren — monatlich, nicht jährlich.')
spacer(doc, 6)

# ══════════════════════════════════════════════════════════════════════════════
# 11. HÄUFIGE FEHLER & WAS TUN WENN...
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '11  Häufige Fehler & Was tun wenn ...')

add_two_col_table(doc, [
    ('Ich habe versehentlich Umsatzsteuer abgeführt (ohne Reverse Charge)',
     'Solange der Bescheid noch nicht bestandskräftig ist (Einspruchsfrist: 1 Monat), kann korrigiert werden. Sofort Steuerberaterin kontaktieren und nachweisen, dass der Umsatz in Irland war.'),
    ('Ich habe noch keine USt-IdNr.',
     'Sofort beim Bundeszentralamt für Steuern in Saarlouis online beantragen. Dauert ca. 1 Woche. Rückwirkende Beantragung ist möglich.'),
    ('Ich bin seit Jahren Kleinunternehmerin',
     'Wechsel zur Umsatzsteuerpflicht ist jederzeit möglich — Zweizeiler ans Finanzamt. Bereits verlorene Vorsteuer ist in den meisten Fällen nicht mehr rückforderbar.'),
    ('Ich habe Steuererklärungen nicht abgegeben',
     'Unbedingt nachreichen! Steuerberater beauftragen. Nicht abgeben ist keine Option — das Finanzamt schätzt dann. Ratenzahlungen können beantragt werden.'),
    ('Mein Steuerberater kennt MONAT / Reverse Charge nicht',
     'Zeig ihm diese Unterlagen und die Provisionsabrechnungen von MONAT. Notfalls spezialisierte Beraterin suchen (z.B. "der Online-Steuerberater" auf Instagram).'),
], header=('Situation', 'Lösung'), col_widths=(5.0, 9.5))

# ══════════════════════════════════════════════════════════════════════════════
# 12. CHECKLISTE
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '12  Deine Checkliste')
add_body(doc, 'Hake Schritt für Schritt ab — so bist du steuerlich auf der sicheren Seite.')
spacer(doc, 4)

add_h2(doc, 'Setup & Anmeldung')
add_checklist_item(doc, 'Gewerbe angemeldet (Haupt- oder Nebenerwerb)')
add_checklist_item(doc, 'Eröffnungsfragebogen beim Finanzamt ausgefüllt')
add_checklist_item(doc, 'Umsatzsteuer-Identifikationsnummer (USt-IdNr.) beantragt')
add_checklist_item(doc, 'Umsatzsteuerpflicht gewählt (empfohlen!) — nicht Kleinunternehmer-Option')
add_checklist_item(doc, 'USt-IdNr. in meinem MONAT-Account hinterlegt')
spacer(doc, 6)

add_h2(doc, 'Laufende Buchhaltung')
add_checklist_item(doc, 'Monatliche Provisionsabrechnungen heruntergeladen und als PDF gespeichert')
add_checklist_item(doc, 'Monatliche Gesamtübersicht (Excel/PDF) erstellt und abgeheftet')
add_checklist_item(doc, 'Alle Betriebsausgaben-Belege gesammelt (digital oder Papier)')
add_checklist_item(doc, 'Bei Bewirtungen: Anlass + Namen auf dem Beleg notiert')
add_checklist_item(doc, 'Monatliche Zusammenfassung an den Steuerberater geschickt')
spacer(doc, 6)

add_h2(doc, 'Steuerliche Meldungen')
add_checklist_item(doc, 'Umsatzsteuervoranmeldung quartalsweise eingereicht')
add_checklist_item(doc, 'Zusammenfassende Meldung (ZM) quartalsweise eingereicht')
add_checklist_item(doc, 'Einkommensteuererklärung jährlich bis 31.07. abgegeben')
add_checklist_item(doc, 'Umsatzsteuerjahreserklärung abgegeben')
spacer(doc, 6)

add_h2(doc, 'Finanzplanung')
add_checklist_item(doc, 'Ca. 1/3 des Gewinns als Steuerrücklage auf separatem Konto')
add_checklist_item(doc, 'Monatlichen Kassensturz gemacht')
add_checklist_item(doc, 'Steuerberaterin gefunden und regelmäßiger Kontakt hergestellt')
spacer(doc, 8)

# ══════════════════════════════════════════════════════════════════════════════
# RECHTLICHER HINWEIS
# ══════════════════════════════════════════════════════════════════════════════
add_hint_box(doc,
    'Rechtlicher Hinweis',
    'Dieser Leitfaden wurde auf Basis eines praxisbezogenen Steuer-Calls erstellt '
    'und dient ausschließlich der allgemeinen Information. Er stellt keine '
    'steuerrechtliche Beratung dar und ersetzt nicht die individuelle Beratung durch '
    'eine qualifizierte Steuerberaterin. Jeder Steuerfall ist individuell — bitte '
    'spreche immer mit deiner Steuerberaterin über deinen persönlichen Sachverhalt.',
    bg='F0F0F0', border_color='999999')

# ── Speichern ──────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Gespeichert: {OUTPUT}")
